"""Document reader tool for extracting text from various file formats."""

from pathlib import Path
from typing import Any

from koda.core.tools.base import Tool


class DocumentReaderTool(Tool):
    """Tool to read and extract text from documents (PDF, Word, TXT, etc.)."""
    
    @property
    def name(self) -> str:
        return "document_reader"
    
    @property
    def description(self) -> str:
        return """Read and extract text content from documents.

Supports:
- PDF files (.pdf)
- Word documents (.doc, .docx)
- Text files (.txt)
- CSV files (.csv)
- JSON files (.json)
- Markdown files (.md)

Use this when the user asks about the content of a file they sent via WhatsApp.
The file path is usually provided in the message metadata or available in the downloads folder.

Examples:
- "What does this document say?" (after user sent a file)
- "Summarize the PDF I just sent"
- "Extract text from ~/downloads/document.docx"
"""
    
    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "The file path to read. For WhatsApp files, check the media_path in the message context."
                }
            },
            "required": ["path"]
        }
    
    async def execute(self, path: str, **kwargs: Any) -> str:
        try:
            file_path = Path(path).expanduser()
            if not file_path.exists():
                return f"❌ File not found: {path}"
            if not file_path.is_file():
                return f"❌ Not a file: {path}"
            
            extension = file_path.suffix.lower()
            
            # Text-based files
            if extension in ['.txt', '.md', '.csv', '.json', '.py', '.js', '.html', '.css', '.xml']:
                try:
                    content = file_path.read_text(encoding="utf-8")
                    return f"📄 **Content of {file_path.name}:**\n\n```\n{content[:8000]}\n```\n\n{'_(truncated...)_' if len(content) > 8000 else ''}"
                except UnicodeDecodeError:
                    # Try binary reading for other encodings
                    return await self._extract_binary_text(file_path)
            
            # PDF files
            elif extension == '.pdf':
                return await self._read_pdf(file_path)
            
            # Word documents
            elif extension in ['.docx', '.doc']:
                return await self._read_word(file_path)
            
            # Image files - describe them
            elif extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
                return f"🖼️ **Image received:** {file_path.name}\n\nSize: {file_path.stat().st_size} bytes\n\n_(Image analysis not available in document reader)_"
            
            # Other binary files
            else:
                return await self._extract_binary_text(file_path)
                
        except PermissionError:
            return f"❌ Permission denied: {path}"
        except Exception as e:
            return f"❌ Error reading file: {str(e)}"
    
    async def _read_pdf(self, file_path: Path) -> str:
        """Extract text from PDF."""
        try:
            # Try PyPDF2 first
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(str(file_path))
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                if text.strip():
                    return f"📄 **PDF Content: {file_path.name}**\n\n```\n{text[:8000]}\n```\n\n{'_(truncated...)_' if len(text) > 8000 else ''}"
                else:
                    return f"📄 **PDF: {file_path.name}**\n\n_(PDF contains no extractable text - may be scanned images)_"
            except ImportError:
                # Fallback to pdfplumber
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        text = ""
                        for page in pdf.pages:
                            text += page.extract_text() or "" + "\n"
                        
                        if text.strip():
                            return f"📄 **PDF Content: {file_path.name}**\n\n```\n{text[:8000]}\n```\n\n{'_(truncated...)_' if len(text) > 8000 else ''}"
                        else:
                            return f"📄 **PDF: {file_path.name}**\n\n_(PDF contains no extractable text)_"
                except ImportError:
                    return f"📄 **PDF: {file_path.name}**\n\nSize: {file_path.stat().st_size} bytes\n\n⚠️ PDF reading libraries not installed. Install with: `pip install PyPDF2` or `pip install pdfplumber`"
        except Exception as e:
            return f"❌ Error reading PDF: {str(e)}"
    
    async def _read_word(self, file_path: Path) -> str:
        """Extract text from Word document."""
        try:
            # Try python-docx for .docx files
            if file_path.suffix.lower() == '.docx':
                try:
                    from docx import Document
                    doc = Document(file_path)
                    text = ""
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                    
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + " "
                            text += "\n"
                    
                    if text.strip():
                        return f"📄 **Word Document: {file_path.name}**\n\n```\n{text[:8000]}\n```\n\n{'_(truncated...)_' if len(text) > 8000 else ''}"
                    else:
                        return f"📄 **Word Document: {file_path.name}**\n\n_(Document appears to be empty or contains no text)_"
                except ImportError:
                    return f"📄 **Word Document: {file_path.name}**\n\nSize: {file_path.stat().st_size} bytes\n\n⚠️ python-docx not installed. Install with: `pip install python-docx`"
            else:
                # .doc files are harder, need antiword or similar
                return f"📄 **Word Document: {file_path.name}**\n\nSize: {file_path.stat().st_size} bytes\n\n⚠️ Old .doc format not supported. Please convert to .docx"
        except Exception as e:
            return f"❌ Error reading Word document: {str(e)}"
    
    async def _extract_binary_text(self, file_path: Path) -> str:
        """Try to extract text from binary files."""
        try:
            # Read as binary and decode what we can
            content = file_path.read_bytes()
            
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text = content.decode(encoding, errors='ignore')
                    # Filter out non-printable characters
                    text = ''.join(char for char in text if char.isprintable() or char in '\n\r\t')
                    text = text.strip()
                    
                    if len(text) > 100:  # If we got meaningful text
                        return f"📄 **Extracted text from {file_path.name}:**\n\n```\n{text[:8000]}\n```\n\n{'_(truncated...)_' if len(text) > 8000 else ''}"
                    break
                except:
                    continue
            
            return f"📎 **Binary file: {file_path.name}**\n\nSize: {file_path.stat().st_size} bytes\nType: {file_path.suffix}\n\n_(This file type cannot be read as text)_"
        except Exception as e:
            return f"❌ Error reading file: {str(e)}"
